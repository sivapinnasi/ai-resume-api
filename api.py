# app.py
from flask import Flask, request, jsonify
from flask_cors import CORS
import requests
import json
import os
import base64
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import logging

app = Flask(__name__)
CORS(app)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Get OpenRouter API key from environment variables
OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY")
if not OPENROUTER_API_KEY:
    logger.warning("OpenRouter API key not found. Set OPENROUTER_API_KEY environment variable.")

# Resume templates
TEMPLATES = {
    "professional": {
        "font": "Calibri",
        "heading_size": 14,
        "subheading_size": 12,
        "body_size": 11,
        "heading_color": "2F5496",
        "use_borders": True
    },
    "creative": {
        "font": "Georgia",
        "heading_size": 16,
        "subheading_size": 13,
        "body_size": 12,
        "heading_color": "7A1712",
        "use_borders": False
    },
    "minimal": {
        "font": "Arial",
        "heading_size": 13,
        "subheading_size": 11,
        "body_size": 10,
        "heading_color": "000000",
        "use_borders": False
    }
}

def generate_resume_content(user_data):
    """
    Generate resume content using OpenRouter API based on job title
    """
    if not OPENROUTER_API_KEY:
        logger.error("OpenRouter API key not available")
        return None
    
    # Prepare prompt for the AI
    prompt = f"""
    Create a professional resume for a {user_data['job_title']} role with the following information:
    
    Skills: {', '.join(user_data['skills'])}
    
    Certificates: {', '.join(user_data['certificates'])}
    
    Projects: {', '.join(user_data['projects'])}
    
    Educational background: {user_data['education']}
    
    Work experience: {user_data['experience']}
    
    Format the resume with appropriate sections for Personal Information, Skills, Professional Experience, Education, Projects, and Certifications.
    For each job in Work Experience, provide 3-4 bullet points highlighting achievements.
    For projects, focus on technologies used and outcomes.
    
    Return the content in JSON format with these sections:
    1. summary - a professional summary tailored to the job title
    2. skills - formatted list of skills grouped by category
    3. experience - formatted work experience with bullet points
    4. education - formatted education details
    5. projects - formatted project descriptions
    6. certificates - formatted certificates list
    """

    try:
        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                "Content-Type": "application/json"
            },
            json={
                "model": "anthropic/claude-3-opus",
                "messages": [
                    {"role": "user", "content": prompt}
                ],
                "response_format": {"type": "json_object"}
            }
        )
        
        if response.status_code == 200:
            result = response.json()
            generated_content = json.loads(result["choices"][0]["message"]["content"])
            return generated_content
        else:
            logger.error(f"OpenRouter API error: {response.status_code} - {response.text}")
            return None
    except Exception as e:
        logger.error(f"Error calling OpenRouter API: {str(e)}")
        return None

def create_docx_resume(user_data, ai_content, template_name):
    """
    Create a docx resume using the chosen template and AI-generated content
    """
    template = TEMPLATES.get(template_name, TEMPLATES["professional"])
    
    doc = Document()
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Personal Information
    heading = doc.add_heading(f"{user_data['name']}", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact_info = doc.add_paragraph()
    contact_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_info.add_run(f"{user_data['email']} | {user_data['phone']} | {user_data['address']}")
    
    # Professional Summary
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph(ai_content['summary'])
    
    # Skills Section
    doc.add_heading('Skills', level=1)
    doc.add_paragraph(ai_content['skills'])
    
    # Experience Section
    doc.add_heading('Professional Experience', level=1)
    doc.add_paragraph(ai_content['experience'])
    
    # Education Section
    doc.add_heading('Education', level=1)
    doc.add_paragraph(ai_content['education'])
    
    # Projects Section
    doc.add_heading('Projects', level=1)
    doc.add_paragraph(ai_content['projects'])
    
    # Certificates Section
    doc.add_heading('Certifications', level=1)
    doc.add_paragraph(ai_content['certificates'])
    
    # Save to bytes
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io.getvalue()

@app.route('/generate-resume', methods=['POST'])
def generate_resume():
    try:
        data = request.json
        
        # Validate required fields
        required_fields = ['name', 'email', 'phone', 'address', 'job_title', 
                           'skills', 'certificates', 'projects', 'education', 
                           'experience', 'template']
        
        for field in required_fields:
            if field not in data:
                return jsonify({"error": f"Missing required field: {field}"}), 400
        
        # Check template validity
        template_name = data['template']
        if template_name not in TEMPLATES:
            return jsonify({"error": f"Invalid template: {template_name}"}), 400
        
        # Generate AI content
        ai_content = generate_resume_content(data)
        if not ai_content:
            return jsonify({"error": "Failed to generate resume content"}), 500
        
        # Create resume document
        resume_bytes = create_docx_resume(data, ai_content, template_name)
        
        # Convert to base64 for sending to frontend
        resume_base64 = base64.b64encode(resume_bytes).decode('utf-8')
        
        return jsonify({
            "success": True,
            "resume": resume_base64,
            "filename": f"{data['name'].replace(' ', '_')}_resume.docx"
        })
        
    except Exception as e:
        logger.error(f"Error generating resume: {str(e)}")
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
